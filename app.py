import json
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from docx import Document
import PyPDF2
import io
import google.generativeai as genai
from dotenv import load_dotenv
import os
import magic
import bleach
from werkzeug.utils import secure_filename
import re
from flask_talisman import Talisman

# Configuración inicial
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB

load_dotenv()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Configuración de CORS con opciones más seguras
CORS(app, resources={
    r"/*": {
        "origins": ["http://localhost:*", "https://yourdomain.com"],
        "methods": ["GET", "POST"],
        "allow_headers": ["Content-Type"]
    }
})

from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# Configuración mejorada de rate limiting
limiter = Limiter(
    app=app,
    key_func=lambda: get_remote_address() or request.headers.get('X-Forwarded-For', get_remote_address()),
    storage_uri="memory://",  # Para producción usa Redis: "redis://localhost:6379"
    strategy="fixed-window",  # Alternativa: "moving-window"
    default_limits=["200 per day", "50 per hour", "10 per minute"]
)

from flask_limiter.errors import RateLimitExceeded

@app.errorhandler(RateLimitExceeded)
def handle_rate_limit_exceeded(e):
    return jsonify({
        "error": "rate_limit_exceeded",
        "message": f"Has excedido el límite de solicitudes. Por favor espera. Límite: {e.description}"
    }), 429

# Configuración de Talisman con políticas de seguridad mejoradas
csp = {
    'default-src': "'self'",
    'script-src': [
        "'self'",
        'https://cdnjs.cloudflare.com',
        'https://cdn.jsdelivr.net',
        "'unsafe-inline'",
        'blob:'
    ],
    'style-src': [
        "'self'",
        'https://cdnjs.cloudflare.com',
        'https://fonts.googleapis.com',
        "'unsafe-inline'"
    ],
    'font-src': [
        "'self'",
        'https://cdnjs.cloudflare.com',
        'https://fonts.gstatic.com'
    ],
    'img-src': [
        "'self'",
        'data:',
        'blob:'
    ],
    'worker-src': [
        "'self'",
        'blob:',
        'https://cdnjs.cloudflare.com'
    ],
    'connect-src': [
        "'self'",
        'blob:',
        'https://cdnjs.cloudflare.com'
    ]
}

Talisman(
    app,
    force_https=False,  # False para desarrollo, True para producción
    strict_transport_security=False,  # False para desarrollo
    content_security_policy=csp,
    session_cookie_secure=False  # False para desarrollo
)

# Configuración de Gemini
gemini_api_key = os.getenv('GEMINI_API_KEY')
genai.configure(api_key=gemini_api_key)


# Middleware para filtrar solicitudes malformadas
@app.before_request
def filter_bad_requests():
    if request.method not in ['GET', 'POST', 'OPTIONS']:
        return jsonify({"error": "Método no permitido"}), 405

    # Rechazar solicitudes sin User-Agent
    if not request.headers.get('User-Agent'):
        return jsonify({"error": "Solicitud no válida"}), 400

    # Validación especial para rutas que no son /process
    if request.method == 'POST' and request.path != '/process' and not request.is_json:
        return jsonify({"error": "Se requiere Content-Type: application/json"}), 415

@app.route('/')
def index():
    return render_template('index.html')

def validate_source(url):
    """Valida que las URLs de fuentes sean seguras y de dominios confiables"""
    if not url:
        return False

    allowed_domains = [
        'scholar.google.com',
        'jstor.org',
        'pubmed.ncbi.nlm.nih.gov',
        'doi.org',
        'arxiv.org',
        'sciencedirect.com',
        'academic.oup.com',
        'springer.com',
        'tandfonline.com',
        'nature.com',
        'science.org',
        'researchgate.net',
        'ncbi.nlm.nih.gov',
        'plos.org',
        'ieee.org',
        'ama-assn.org'
    ]

    try:
        # Verificar formato URL
        if not re.match(r'^https?://[^\s/$.?#].[^\s]*$', url, re.IGNORECASE):
            return False

        # Extraer dominio
        domain = re.search(
            r'https?://(?:www\.)?([^/]+)',
            url.lower()
        ).group(1)

        # Eliminar subdominios no relevantes (ej: 'www')
        domain_parts = domain.split('.')
        if len(domain_parts) > 2:
            domain = '.'.join(domain_parts[-2:])

        # Verificar contra dominios permitidos
        return any(
            allowed == domain or domain.endswith(f'.{allowed}')
            for allowed in allowed_domains
        )

    except:
        return False


def allowed_file(filename):
    """Verifica que la extensión sea válida"""
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def validate_file_type(file_stream, filename):
    """Valida el tipo MIME real del archivo"""
    mime = magic.Magic(mime=True)
    file_mime = mime.from_buffer(file_stream.read(2048))
    file_stream.seek(0)  # Rebobinar para no afectar el procesamiento posterior

    valid_mimes = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }

    ext = os.path.splitext(filename)[1].lower()
    return file_mime == valid_mimes.get(ext)


def sanitize_text(text):
    if not text:
        return ""

    # Limpiar HTML y caracteres peligrosos
    cleaned = bleach.clean(text,
                           tags=[],
                           attributes={},
                           strip=True)

    # Eliminar caracteres no imprimibles excepto saltos de línea
    cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned)

    # Limitar longitud
    return cleaned[:50000]

def validate_pdf(file_stream):
    try:
        PyPDF2.PdfReader(file_stream)
        return True
    except PyPDF2.PdfReadError:
        return False

def validate_docx(file_stream):
    try:
        Document(file_stream)
        return True
    except:
        return False


def extract_text(file):
    if file.filename.endswith('.docx'):
        doc = Document(io.BytesIO(file.read()))
        return [{"text": p.text} for p in doc.paragraphs if p.text.strip()]
    elif file.filename.endswith('.pdf'):
        pdf = PyPDF2.PdfReader(file)
        text = []
        for page in pdf.pages:
            text.append(page.extract_text())
        return [{"text": t} for t in text if t.strip()]
    return []


@app.route('/chat', methods=['POST'])
@limiter.limit("10 per minute")
def chat_with_document():
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        document_text = data.get('document_text', '').strip()

        if not document_text:
            return jsonify({"error": "Texto del documento vacío"}), 400

        # Sanitizar entradas
        question = sanitize_text(question)
        if question == "":
            return jsonify({
                "answer": 'Ingrese una pregunta válida',
                "external_source": None  # None si no es válida
            }), 400
        document_text = sanitize_text(document_text)

        model = genai.GenerativeModel('gemini-1.5-flash')

        prompt = (
            "Actúa como un asistente académico especializado en analizar documentos. "
            "Responde la pregunta del usuario basándote PRINCIPALMENTE en el contenido "
            "del documento proporcionado. Sigue estas reglas:\n\n"

            "1) SI la respuesta está DIRECTAMENTE en el documento:\n"
            "- Extrae la información precisa\n"
            "- Cita el fragmento relevante entre comillas \"\"\n"
            "- Responde de manera concisa\n\n"

            "2) SI necesitas COMPLEMENTAR con información externa:\n"
            "- Primero indica claramente \"Según el documento:\"\n"
            "- Luego añade \"Información adicional:\" con datos relevantes\n"
            "- Proporciona EXACTAMENTE 1 referencia académica confiable\n\n"

            "3) SI la pregunta NO está relacionada con el documento:\n"
            "- Responde amablemente que la respuesta a esa pregunta no se encuentra en el documento subido\n"
            "- Responde la pregunta de manera concisa añadiendo que según (inserte la fuente aquí) \n\n"

            f"DOCUMENTO:\n{document_text}\n\n"
            f"PREGUNTA DEL USUARIO:\n{question}"
        )

        response = model.generate_content(prompt)

        # Procesar la respuesta para identificar fuentes externas
        answer = response.text
        external_source = None

        if "Información adicional:" in answer:
            source_match = re.search(r'Fuente:\s*(.+?)\s*(\(https?://[^\s]+)?', answer)
            if source_match:
                url = source_match.group(2)[1:-1] if source_match.group(2) else None
                if url and validate_source(url):  # <- VALIDACIÓN DE URL
                    external_source = {
                        "name": sanitize_text(source_match.group(1))[:200],
                        "url": url[:500]
                    }

        return jsonify({
            "answer": answer,
            "external_source": external_source  # None si no es válida
        })

    except Exception as e:
        app.logger.error(f"Error en chat: {str(e)}")
        return jsonify({"error": "Error en el servidor"}), 500


@app.route('/suggestions', methods=['POST'])
def generate_questions():
    try:
        data = request.get_json()
        document_text = data.get('document_text', '').strip()

        if not document_text:
            return jsonify({"error": "Texto del documento vacío"}), 400

        model = genai.GenerativeModel('gemini-1.5-flash')

        response = model.generate_content(
            "Genera exactamente 5 preguntas frecuentes breves (máximo 15 palabras cada una) "
            "basadas en este documento. Devuélvelas como una lista JSON:\n\n"
            f"{document_text}"
        )

        # Extraer las preguntas de la respuesta
        questions = []
        print(response.text)
        if response.text.startswith('[') and response.text.endswith(']'):
            try:
                questions = json.loads(response.text)
            except:
                # Si falla el parseo, intentar extraer preguntas de otro formato
                questions = [q.strip() for q in response.text.split('\n') if q.strip()]
        else:
            questions = [q.strip() for q in response.text.split('\n') if q.strip()]

        return jsonify({"questions": questions[2:7]})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/process', methods=['POST'])
@limiter.limit("5 per minute")
def process_file():
    try:
        # Verificar si se envió el archivo correctamente
        if 'file' not in request.files:
            app.logger.error("No se encontró el campo 'file' en la solicitud")
            return jsonify({"error": "No se proporcionó archivo"}), 400

        file = request.files['file']

        # Verificar nombre de archivo
        if file.filename == '':
            app.logger.error("Nombre de archivo vacío")
            return jsonify({"error": "Nombre de archivo vacío"}), 400

        # Validar extensión
        if not allowed_file(file.filename):
            app.logger.error(f"Tipo de archivo no permitido: {file.filename}")
            return jsonify({"error": "Solo se permiten archivos PDF o DOCX"}), 400

        # Validar tipo MIME real
        mime = magic.Magic(mime=True)
        file_stream = file.stream.read(2048)
        file_mime = mime.from_buffer(file_stream)
        file.stream.seek(0)  # Rebobinar para procesar después

        valid_mimes = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }

        ext = os.path.splitext(file.filename)[1].lower()
        if file_mime != valid_mimes.get(ext):
            app.logger.error(f"MIME type no coincide: {file_mime} para extensión {ext}")
            return jsonify({"error": "Tipo de archivo no válido"}), 400

        # Validar estructura del archivo
        if ext == '.pdf' and not validate_pdf(file.stream):
            file.stream.seek(0)
            return jsonify({"error": "El archivo PDF está corrupto"}), 400
        elif ext == '.docx' and not validate_docx(file.stream):
            file.stream.seek(0)
            return jsonify({"error": "El archivo DOCX está corrupto"}), 400

        file.stream.seek(0)
        paragraphs = extract_text(file)

        # Asegurarse de devolver un array incluso para PDFs
        if file.filename.endswith('.pdf'):
            # Para PDFs, convertimos el texto en párrafos
            full_text = ' '.join([p['text'] for p in paragraphs if p['text']])
            paragraphs = [{'text': p} for p in full_text.split('\n\n') if p.strip()]

        return jsonify(paragraphs)

    except Exception as e:
        app.logger.error(f"Error al procesar archivo: {str(e)}")
        return jsonify({"error": "Error al procesar el archivo"}), 500


@app.route('/complement', methods=['POST'])
@limiter.limit("5 per minute")
def complement_info():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()

        if not text:
            return jsonify({"error": "Texto vacío"}), 400

        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            f"Como experto académico, analiza el texto proporcionado y complementa su información con:\n"

            f"1) Datos adicionales relevantes (contexto teórico, cifras actualizadas, "
            f"ejemplos prácticos o controversias académicas).\n"
            
            f"EXACTAMENTE 2 referencias académicas confiables y especializadas en formato:\n"
            f"• NOMBRE_FUENTE (URL_OFICIAL)\n"
            f"(Prioriza fuentes institucionales, revistas científicas o bases de datos reconocidas "
            f"como PubMed, JSTOR, o repositorios " f"universitarios. Evita blogs o sitios sin revisión por pares).\n"
            
            f"Asegúrate de que las referencias respalden directamente los datos agregados y "
            f"estén vinculadas al tema central del texto.\n"
            f"Texto para analizar:\n"
            f"{text}"
        )

        if not response.text:
            return jsonify({"error": "Respuesta inválida"}), 500

        content = response.text.replace('•', '•')
        raw_sources = re.findall(
            r'(?:•|\d+\.)\s*([^\(\n]+?)\s*\((\bhttps?:\/\/[^\s\)]+)\)',
            content,
            flags=re.IGNORECASE
        )

        # Filtrar fuentes válidas
        valid_sources = []
        for name, url in raw_sources[:2]:  # Limitar a 2 fuentes como antes
            if validate_source(url):  # <- AQUÍ USAMOS LA VALIDACIÓN
                clean_name = re.sub(r'[\*\#]', '', name).strip()
                valid_sources.append({
                    "name": clean_name[:200],  # Limitar longitud del nombre
                    "url": url[:500]  # Limitar longitud de URL
                })

        # Si no hay fuentes válidas, usar fuente por defecto
        if not valid_sources:
            valid_sources = [{
                "name": "Google Scholar",
                "url": "https://scholar.google.com"
            }]

        return jsonify({
            "complement": content,
            "sources": valid_sources  # Fuentes validadas
        })

    except Exception as e:
        app.logger.error(f"Error: {str(e)}")
        return jsonify({"error": "Error en el servidor"}), 500

if __name__ == "__main__":
    # Configuración para desarrollo vs producción
    if os.environ.get('FLASK_ENV') == 'production':
        # Configuración para producción
        app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), ssl_context='adhoc')
    else:
        # Configuración para desarrollo
        app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)